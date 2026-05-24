from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path(__file__).resolve().parent
OUT = BASE / "готовые документы"
OUT.mkdir(exist_ok=True)
LOGO = OUT / "логотип ннгу_image1.png"


NNGU = {
    "full": 'Федеральное государственное автономное образовательное учреждение высшего образования\n'
    '"Национальный исследовательский Нижегородский государственный университет им. Н.И. Лобачевского"',
    "short": "ННГУ им. Н.И. Лобачевского",
    "addr": "603022, г. Нижний Новгород, проспект Гагарина, д. 23",
    "inn": "ИНН/КПП 5262004442/526201001",
    "ogrn": "ОГРН 1025203733510",
    "web": "www.unn.ru, e-mail: rector@unn.ru",
}

YANDEX = {
    "full": 'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ "ЯНДЕКС"',
    "short": 'ООО "ЯНДЕКС"',
    "addr": "119021, г. Москва, ул. Льва Толстого, д. 16",
    "inn": "ИНН/КПП 7736207543/770401001",
    "ogrn": "ОГРН 1027700229193",
    "mail": "e-mail: register-yandex@yandex-team.ru",
}

STUDENT = "Токарева Алена Дмитриевна"
STUDENT_ACC = "Токареву Алену Дмитриевну"
GROUP = "3523Б3ПИ1"
INSTITUTE = "Институт информационных технологий, математики и механики"
INSTITUTE_GEN = "Института информационных технологий, математики и механики"
DIRECTION = "09.03.03 Прикладная информатика"
PRACTICE_NOM = "производственная практика (технологическая (проектно-технологическая) практика)"
PRACTICE_GEN = "производственной практики (технологической (проектно-технологической) практики)"
PERIOD = "с 08.06.2026 по 05.07.2026"
NNGU_SUPERVISOR = "доцент кафедры программной инженерии, к.т.н. Кузнецова Мария Викторовна"
ORG_SUPERVISOR = 'ведущий разработчик отдела образовательных проектов ООО "ЯНДЕКС" Смирнов Андрей Павлович'


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:{}".format(key)), str(value))


def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "FFFFFF")
        tbl_borders.append(element)
    tbl_pr.append(tbl_borders)
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                left={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                right={"val": "none", "sz": "0", "color": "FFFFFF"},
                insideH={"val": "none", "sz": "0", "color": "FFFFFF"},
                insideV={"val": "none", "sz": "0", "color": "FFFFFF"},
            )


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def base_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.1
    styles["Normal"].paragraph_format.space_after = Pt(0)
    return doc


def add_para(doc, text="", align=None, bold=False, size=12, first_line=False, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    return p


def add_signature(doc, title, name):
    p = add_para(doc)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(f"{title}\tПодпись\t{name}")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def add_date_number_line(doc, date_text, number_text, reply=False):
    p = add_para(doc)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(5.2), WD_TAB_ALIGNMENT.LEFT)
    p.add_run(f"{date_text}\t№ {number_text}")
    if reply:
        p = add_para(doc)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(5.2), WD_TAB_ALIGNMENT.LEFT)
        p.add_run("На № __________\tот __________")


def add_yandex_corner_blank(doc, include_reply=True, filled_date=None, filled_number=None, reply_to=None):
    lines = [
        YANDEX["full"],
        f"({YANDEX['short']})",
        YANDEX["addr"],
        YANDEX["inn"],
        YANDEX["ogrn"],
        YANDEX["mail"],
    ]
    for i, line in enumerate(lines):
        add_para(doc, line, bold=i == 0, size=9 if i > 1 else 10)
    doc.add_paragraph()
    date = filled_date or "__________"
    num = filled_number or "__________"
    add_date_number_line(doc, date, num, reply=False)
    if include_reply:
        p = add_para(doc)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(5.2), WD_TAB_ALIGNMENT.LEFT)
        if reply_to:
            p.add_run(f"На № {reply_to[0]}\tот {reply_to[1]}")
        else:
            p.add_run("На № __________\tот __________")


def add_nngu_corner_blank(doc, filled_date=None, filled_number=None):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(str(LOGO), width=Cm(1.0))
    for i, line in enumerate(
        [
            "МИНОБРНАУКИ РОССИИ",
            NNGU["full"],
            f"({NNGU['short']})",
            NNGU["addr"],
            NNGU["inn"],
            NNGU["ogrn"],
            NNGU["web"],
        ]
    ):
        add_para(doc, line, bold=i in (0, 1), size=8.5 if i > 2 else 9.5)
    doc.add_paragraph()
    add_date_number_line(doc, filled_date or "__________", filled_number or "__________")


def add_nngu_order_blank(doc, date_text, number_text):
    p = add_para(doc, "МИНОБРНАУКИ РОССИИ", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)
    p.paragraph_format.space_after = Pt(2)
    add_para(doc, NNGU["full"], WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=10)
    add_para(doc, f"({NNGU['short']})", WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_para(doc, "ПРИКАЗ", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=8)
    p = add_para(doc)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(7.8), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(f"{date_text}\t№ {number_text}\tг. Нижний Новгород")


def add_yandex_order_blank(doc, date_text=None, number_text=None):
    add_para(doc, YANDEX["full"], WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
    add_para(doc, f"({YANDEX['short']})", WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, YANDEX["addr"], WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_para(doc, f"{YANDEX['inn']}; {YANDEX['ogrn']}", WD_ALIGN_PARAGRAPH.CENTER, size=10)
    add_para(doc, "ПРИКАЗ", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, space_after=8)
    p = add_para(doc)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(7.8), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(f"{date_text or '__________'}\t№ {number_text or '__________'}\tг. Москва")


def add_address_block(doc, lines):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(8.5)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)


def add_student_table(doc):
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    widths = [Cm(1.2), Cm(5.2), Cm(3.0), Cm(6.0)]
    headers = ["№ п/п", "ФИО студента", "Группа", "Направление подготовки"]
    for col, width in enumerate(widths):
        table.columns[col].width = width
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.cell(0, i), "EDEDED")
    values = ["1", STUDENT, GROUP, DIRECTION]
    for i, value in enumerate(values):
        set_cell_text(table.cell(1, i), value, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT)
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top={"val": "single", "sz": "6", "color": "000000"},
                left={"val": "single", "sz": "6", "color": "000000"},
                bottom={"val": "single", "sz": "6", "color": "000000"},
                right={"val": "single", "sz": "6", "color": "000000"},
            )


def save(doc, name):
    path = OUT / name
    doc.save(path)
    print(path)


def make_yandex_letter_blank():
    doc = base_doc()
    add_yandex_corner_blank(doc)
    add_para(doc, "\n\n\n\n", WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "Заголовок к тексту письма", WD_ALIGN_PARAGRAPH.LEFT, size=12)
    save(doc, "01_бланк_письма_ООО_Яндекс.docx")


def make_yandex_order_blank():
    doc = base_doc()
    add_yandex_order_blank(doc)
    add_para(doc, "\nО заголовке приказа", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
    add_para(doc, "\nПРИКАЗЫВАЮ:", bold=True)
    save(doc, "02_бланк_приказа_ООО_Яндекс.docx")


def make_nngu_letter():
    doc = base_doc()
    add_nngu_corner_blank(doc, "19.05.2026", "02-16/1245")
    add_address_block(
        doc,
        [
            'Генеральному директору ООО "ЯНДЕКС"',
            "А.Г. Савиновскому",
            YANDEX["addr"],
        ],
    )
    add_para(doc, "О направлении студента на практику", bold=True, space_after=8)
    add_para(doc, "Уважаемый Артем Геннадьевич!", WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(
        doc,
        f"{NNGU['short']} просит принять для прохождения {PRACTICE_GEN} студентку 3 курса заочной формы обучения "
        f"{INSTITUTE_GEN}, обучающуюся по направлению подготовки бакалавриата {DIRECTION}:",
        first_line=True,
    )
    add_student_table(doc)
    add_para(doc, f"Сроки прохождения практики: {PERIOD}.", first_line=True, space_after=4)
    add_para(doc, f"Руководитель практики от ННГУ: {NNGU_SUPERVISOR}.", first_line=True)
    add_para(
        doc,
        "Просим назначить руководителя практики от профильной организации и подтвердить возможность прохождения практики ответным письмом.",
        first_line=True,
        space_after=12,
    )
    add_signature(doc, "Директор ИИТММ", "С.В. Козлов")
    save(doc, "03_письмо_ННГУ_о_направлении_на_практику.docx")


def make_yandex_response():
    doc = base_doc()
    add_yandex_corner_blank(doc, include_reply=True, filled_date="22.05.2026", filled_number="17/05-26", reply_to=("02-16/1245", "19.05.2026"))
    add_address_block(
        doc,
        [
            f"Директору {INSTITUTE_GEN}",
            "С.В. Козлову",
            NNGU["addr"],
        ],
    )
    add_para(doc, "О согласии принять студента на практику", bold=True, space_after=8)
    add_para(doc, "Уважаемый Сергей Владимирович!", WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(
        doc,
        f'ООО "ЯНДЕКС" подтверждает готовность принять {STUDENT_ACC} для прохождения {PRACTICE_GEN} {PERIOD}.',
        first_line=True,
    )
    add_para(
        doc,
        f"Руководителем практики от профильной организации назначен {ORG_SUPERVISOR}.",
        first_line=True,
    )
    add_para(
        doc,
        "Организация обеспечит проведение вводного инструктажа, предоставление рабочего места и материалов, необходимых для выполнения индивидуального задания.",
        first_line=True,
        space_after=12,
    )
    add_signature(doc, "Директор по персоналу", "М.А. Иванова")
    save(doc, "04_ответное_письмо_ООО_Яндекс.docx")


def make_nngu_order():
    doc = base_doc()
    add_nngu_order_blank(doc, "05.06.2026", "2154-ОП")
    add_para(doc, "\nО направлении студента на производственную практику", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=8)
    add_para(
        doc,
        f"В соответствии с учебным планом по направлению подготовки {DIRECTION}, Положением о практической подготовке обучающихся ННГУ и на основании письма ООО \"ЯНДЕКС\" от 22.05.2026 № 17/05-26",
        first_line=True,
    )
    add_para(doc, "ПРИКАЗЫВАЮ:", bold=True, space_after=4)
    items = [
        f"Направить {STUDENT_ACC}, студентку 3 курса заочной формы обучения группы {GROUP}, для прохождения {PRACTICE_GEN} в ООО \"ЯНДЕКС\" {PERIOD}.",
        f"Назначить руководителем практики от ННГУ: {NNGU_SUPERVISOR}.",
        f"Согласовать руководителя практики от профильной организации: {ORG_SUPERVISOR}.",
        "Студентке представить отчетные документы по практике в установленный образовательной программой срок.",
        "Контроль за исполнением настоящего приказа возложить на директора ИИТММ С.В. Козлова.",
    ]
    for i, item in enumerate(items, 1):
        add_para(doc, f"{i}. {item}", first_line=True)
    add_para(doc, "Основание: письмо ООО \"ЯНДЕКС\" от 22.05.2026 № 17/05-26.", first_line=True, space_after=12)
    add_signature(doc, "Проректор по образовательной деятельности", "Е.Е. Черных")
    save(doc, "05_приказ_ННГУ_о_направлении_на_практику.docx")


def make_yandex_order():
    doc = base_doc()
    add_yandex_order_blank(doc, "05.06.2026", "Я-125/26")
    add_para(doc, "\nО назначении руководителя практики", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, space_after=8)
    add_para(
        doc,
        f"В целях организации прохождения {PRACTICE_GEN} студенткой {NNGU['short']} {STUDENT} {PERIOD}",
        first_line=True,
    )
    add_para(doc, "ПРИКАЗЫВАЮ:", bold=True, space_after=4)
    items = [
        f"Принять {STUDENT_ACC}, студентку группы {GROUP}, для прохождения практики в ООО \"ЯНДЕКС\" {PERIOD}.",
        f"Назначить руководителем практики от профильной организации: {ORG_SUPERVISOR}.",
        "Руководителю практики провести вводный инструктаж, определить перечень выполняемых заданий и обеспечить контроль соблюдения правил внутреннего трудового распорядка.",
        "По окончании практики подготовить характеристику и подтвердить выполнение программы практики.",
        "Контроль за исполнением настоящего приказа оставляю за собой.",
    ]
    for i, item in enumerate(items, 1):
        add_para(doc, f"{i}. {item}", first_line=True)
    add_para(doc, "Основание: письмо ННГУ им. Н.И. Лобачевского от 19.05.2026 № 02-16/1245.", first_line=True, space_after=12)
    add_signature(doc, "Генеральный директор", "А.Г. Савиновский")
    save(doc, "06_приказ_ООО_Яндекс_о_назначении_руководителя_практики.docx")


if __name__ == "__main__":
    make_yandex_letter_blank()
    make_yandex_order_blank()
    make_nngu_letter()
    make_yandex_response()
    make_nngu_order()
    make_yandex_order()
