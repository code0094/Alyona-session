from pathlib import Path
import re
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SQL = ROOT / "sql"
OUT = DOCS / "coursework-report.docx"
ER_IMAGE = DOCS / "er-chen.png"
REL_IMAGE = DOCS / "relational-schema.png"


def font(size: int, bold: bool = False):
    name = "arialbd.ttf" if bold else "arial.ttf"
    return ImageFont.truetype(str(Path("C:/Windows/Fonts") / name), size=size)


def center_text(draw, box, text, fnt, fill=(20, 32, 44)):
    x1, y1, x2, y2 = box
    lines = []
    for raw in text.split("\n"):
        lines.extend(textwrap.wrap(raw, width=max(10, int((x2 - x1) / (fnt.size * 0.62)))))
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total_h = sum(heights) + 6 * (len(lines) - 1)
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, h in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        x = x1 + ((x2 - x1) - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, font=fnt, fill=fill)
        y += h + 6


def rectangle(draw, box, text):
    draw.rounded_rectangle(box, radius=10, fill=(235, 244, 255), outline=(25, 85, 150), width=3)
    center_text(draw, box, text, font(24, True))


def ellipse(draw, box, text):
    draw.ellipse(box, fill=(255, 250, 230), outline=(160, 120, 30), width=3)
    center_text(draw, box, text, font(16))


def diamond(draw, center, w, h, text):
    cx, cy = center
    points = [(cx, cy - h // 2), (cx + w // 2, cy), (cx, cy + h // 2), (cx - w // 2, cy)]
    draw.polygon(points, fill=(238, 255, 238), outline=(50, 130, 70))
    draw.line(points + [points[0]], fill=(50, 130, 70), width=3)
    center_text(draw, (cx - w // 2, cy - h // 2, cx + w // 2, cy + h // 2), text, font(19, True))


def draw_line(draw, start, end, label=None):
    draw.line((start, end), fill=(60, 70, 80), width=3)
    if label:
        x = (start[0] + end[0]) // 2
        y = (start[1] + end[1]) // 2
        draw.text((x + 6, y - 22), label, font=font(18, True), fill=(70, 70, 70))


def oval_anchor(box, target):
    x1, y1, x2, y2 = box
    cx = (x1 + x2) // 2
    cy = (y1 + y2) // 2
    if target[0] < x1:
        return (x1, cy)
    if target[0] > x2:
        return (x2, cy)
    if target[1] < y1:
        return (cx, y1)
    return (cx, y2)


def make_er_image():
    img = Image.new("RGB", (2200, 1450), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 25), "ER-диаграмма в нотации Чена", font=font(34, True), fill=(20, 32, 44))

    type_box = (130, 260, 430, 390)
    laptop_box = (910, 230, 1240, 430)
    supplier_box = (130, 860, 430, 990)
    delivery_box = (910, 830, 1240, 1010)
    rectangle(draw, type_box, "Тип\nноутбука")
    rectangle(draw, laptop_box, "Ноутбук")
    rectangle(draw, supplier_box, "Поставщик")
    rectangle(draw, delivery_box, "Поставка")

    diamond(draw, (650, 325), 190, 110, "относится")
    diamond(draw, (650, 925), 190, 110, "выполняет")
    diamond(draw, (1075, 620), 190, 110, "содержит")

    draw_line(draw, (430, 325), (555, 325), "1")
    draw_line(draw, (745, 325), (910, 325), "N")
    draw_line(draw, (430, 925), (555, 925), "1")
    draw_line(draw, (745, 925), (910, 925), "N")
    draw_line(draw, (1075, 430), (1075, 565), "1")
    draw_line(draw, (1075, 675), (1075, 830), "N")

    attrs = [
        ((95, 130, 335, 205), "type_id\nPK", (260, 260)),
        ((360, 130, 620, 205), "type_name", (360, 260)),
        ((1380, 95, 1620, 170), "laptop_id\nPK", (1160, 230)),
        ((1650, 95, 1870, 170), "name", (1240, 245)),
        ((1900, 150, 2120, 225), "model", (1240, 265)),
        ((1360, 230, 1580, 305), "width_m", (1240, 285)),
        ((1610, 270, 1830, 345), "depth_m", (1240, 305)),
        ((1870, 325, 2090, 400), "height_m", (1240, 325)),
        ((1360, 400, 1580, 475), "processor", (1240, 345)),
        ((1610, 455, 1830, 530), "ram_mb", (1240, 365)),
        ((1870, 515, 2090, 590), "hdd_gb", (1240, 385)),
        ((1360, 585, 1580, 660), "drive_type", (1240, 405)),
        ((1610, 650, 1870, 725), "monitor_inches", (1200, 430)),
        ((1880, 710, 2160, 785), "operating_system", (1160, 430)),
        ((1390, 765, 1610, 840), "unit_cost", (1020, 430)),
        ((1390, 820, 1630, 895), "delivery_id\nPK", (1240, 860)),
        ((1660, 890, 1900, 965), "delivery_date", (1240, 900)),
        ((1390, 995, 1610, 1070), "quantity", (1240, 940)),
        ((1660, 1080, 1880, 1155), "unit_price", (1240, 980)),
        ((95, 700, 335, 775), "supplier_id\nPK", (260, 860)),
        ((360, 1110, 640, 1185), "supplier_name", (430, 930)),
        ((95, 1210, 315, 1285), "city", (230, 990)),
        ((350, 1240, 570, 1315), "address", (310, 990)),
        ((110, 1325, 330, 1400), "phone", (250, 990)),
    ]
    for box, text, target in attrs:
        ellipse(draw, box, text)
        draw_line(draw, oval_anchor(box, target), target)

    img.save(ER_IMAGE)


def make_rel_image():
    img = Image.new("RGB", (1700, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.text((40, 25), "Реляционная схема базы данных", font=font(34, True), fill=(20, 32, 44))

    tables = [
        ((70, 150, 450, 360), "laptop_types", ["PK type_id", "type_name", "description"]),
        ((650, 120, 1110, 520), "laptops", ["PK laptop_id", "FK type_id", "name", "model", "width_m, depth_m, height_m", "processor, ram_mb, hdd_gb", "drive_type, monitor_inches", "operating_system", "unit_cost"]),
        ((70, 650, 450, 900), "suppliers", ["PK supplier_id", "supplier_name", "city", "address", "phone"]),
        ((650, 650, 1110, 940), "deliveries", ["PK delivery_id", "FK supplier_id", "FK laptop_id", "delivery_date", "quantity", "unit_price"]),
    ]
    for box, title, rows in tables:
        draw.rounded_rectangle(box, radius=10, fill=(246, 248, 250), outline=(80, 90, 100), width=3)
        draw.rectangle((box[0], box[1], box[2], box[1] + 48), fill=(25, 85, 150))
        center_text(draw, (box[0], box[1], box[2], box[1] + 48), title, font(22, True), fill="white")
        y = box[1] + 62
        for row in rows:
            draw.text((box[0] + 18, y), row, font=font(17), fill=(20, 32, 44))
            y += 34
    draw_line(draw, (450, 250), (650, 250), "1:N")
    draw_line(draw, (450, 770), (650, 770), "1:N")
    draw_line(draw, (880, 520), (880, 650), "1:N")
    img.save(REL_IMAGE)


def read(path):
    return path.read_text(encoding="utf-8")


def split_queries(sql_text):
    parts = []
    current_title = None
    current_lines = []
    for line in sql_text.splitlines():
        if line.startswith("-- "):
            if current_title and current_lines:
                parts.append((current_title, "\n".join(current_lines).strip()))
                current_lines = []
            current_title = line[3:].strip()
        elif current_title:
            current_lines.append(line)
    if current_title and current_lines:
        parts.append((current_title, "\n".join(current_lines).strip()))
    return parts


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def build_markdown(schema, seed, queries):
    md = [
        "# Курсовая работа по дисциплине «Базы данных»",
        "",
        "Тема: база данных и приложение для учета поставок ноутбуков в салон компьютерной техники.",
        "",
        "Студентка: Токарева. Вариант: 11.",
        "",
        "## Постановка задачи",
        "",
        "Салон компьютерной техники получает ноутбуки от нескольких производителей. Необходимо создать базу данных, которая хранит типы ноутбуков, технические характеристики ноутбуков, сведения о поставщиках и данные о поставках. На основе базы требуется выполнить запросы варианта 11 и разработать простое приложение для чтения и записи данных на сервере MySQL.",
        "",
        "## Реляционная модель",
        "",
        "![Реляционная схема](relational-schema.png)",
        "",
        "- laptop_types(type_id PK, type_name, description)",
        "- laptops(laptop_id PK, type_id FK, name, model, width_m, depth_m, height_m, processor, ram_mb, hdd_gb, drive_type, monitor_inches, operating_system, unit_cost)",
        "- suppliers(supplier_id PK, supplier_name, city, address, phone)",
        "- deliveries(delivery_id PK, supplier_id FK, laptop_id FK, delivery_date, quantity, unit_price)",
        "",
        "Поле `unit_cost` хранит себестоимость ноутбука. В тексте варианта есть примечание, что прибыль считается как цена продажи минус себестоимость, поэтому себестоимость нужна для корректной экономической интерпретации поставок.",
        "",
        "## ER-диаграмма в нотации Чена",
        "",
        "![ER-диаграмма в нотации Чена](er-chen.png)",
        "",
        "## CREATE TABLE",
        "",
        "```sql",
        schema.strip(),
        "```",
        "",
        "## Тестовые данные",
        "",
        "```sql",
        seed.strip(),
        "```",
        "",
        "## SQL-запросы по варианту",
        "",
    ]
    for title, sql in split_queries(queries):
        md += [f"### {title}", "", "```sql", sql, "```", ""]
    md += [
        "## Приложение",
        "",
        "Приложение выполнено на C# WinForms с использованием MySqlConnector. Оно подключается к MySQL, показывает таблицы ноутбуков, поставщиков и поставок, выполняет основные SELECT-запросы варианта, добавляет и удаляет записи поставок, выполняет UPDATE для Rover1000 и DELETE для SComp.",
        "",
        "## Вывод",
        "",
        "В работе спроектирована база данных для учета поставок ноутбуков, построены ER-диаграмма и реляционная модель, подготовлены SQL-скрипты создания и заполнения таблиц, выполнены запросы варианта 11. Разработанное приложение подтверждает возможность чтения и изменения данных на сервере MySQL.",
    ]
    (DOCS / "coursework-report.md").write_text("\n".join(md) + "\n", encoding="utf-8")


def add_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = rows[0][i]
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
    return table


def build_docx(schema, seed, queries):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Курсовая работа\nпо дисциплине «Базы данных»")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\nТема: разработка базы данных и приложения для учета поставок ноутбуков в салон компьютерной техники\n").bold = True
    meta.add_run("Студентка: Токарева\nФорма обучения: заочная\nВариант: 11\n")
    doc.add_page_break()

    doc.add_heading("Постановка задачи", level=1)
    doc.add_paragraph(
        "Салон компьютерной техники получает ноутбуки от нескольких производителей, "
        "которые предлагают типы ноутбуков: Nautilus, экономическая модель, мини, "
        "мобильная модель, производительная модель. Нужно создать базу данных, "
        "включающую сведения о ноутбуках, поставщиках и поставках, выполнить SQL-запросы "
        "варианта 11 и разработать простое приложение для чтения и записи информации с сервера MySQL."
    )
    doc.add_paragraph(
        "Поле себестоимости unit_cost включено в таблицу ноутбуков, потому что в варианте указано, "
        "что прибыль определяется как разница между ценой продажи и себестоимостью. "
        "Это поле позволяет расширить запросы экономической интерпретацией поставок."
    )

    doc.add_heading("ER-диаграмма в нотации Чена", level=1)
    doc.add_paragraph("На диаграмме прямоугольниками показаны сущности, ромбами - связи, овалами - основные атрибуты.")
    doc.add_picture(str(ER_IMAGE), width=Inches(6.4))

    doc.add_heading("Реляционная модель", level=1)
    doc.add_picture(str(REL_IMAGE), width=Inches(6.4))
    add_table(doc, [
        ["Таблица", "Первичный ключ", "Внешние ключи", "Назначение"],
        ["laptop_types", "type_id", "-", "Справочник типов ноутбуков"],
        ["laptops", "laptop_id", "type_id -> laptop_types.type_id", "Ноутбуки и характеристики"],
        ["suppliers", "supplier_id", "-", "Поставщики"],
        ["deliveries", "delivery_id", "supplier_id -> suppliers.supplier_id\nlaptop_id -> laptops.laptop_id", "Партии поставок"],
    ], widths=[1.2, 1.25, 2.35, 1.6])

    doc.add_heading("Запросы на создание таблиц", level=1)
    add_code(doc, schema.strip())

    doc.add_heading("Тестовые данные", level=1)
    add_code(doc, seed.strip())

    doc.add_heading("SQL-запросы по варианту 11", level=1)
    for title, sql in split_queries(queries):
        doc.add_heading(title, level=2)
        add_code(doc, sql)

    doc.add_heading("Описание приложения", level=1)
    doc.add_paragraph(
        "Приложение создано на C# WinForms. Для подключения к MySQL используется пакет MySqlConnector. "
        "Пользователь может указать параметры подключения, просмотреть таблицы ноутбуков, поставщиков и поставок, "
        "выполнить учебные SELECT-запросы, добавить новую поставку, удалить выбранную поставку, "
        "изменить объем оперативной памяти Rover1000 и удалить поставщика SComp. "
        "Таким образом приложение демонстрирует получение и запись информации с сервера."
    )
    add_table(doc, [
        ["Функция", "Назначение"],
        ["Подключение к MySQL", "Проверка доступа к серверу и базе coursework_db"],
        ["Просмотр таблиц", "Отображение данных ноутбуков, поставщиков и поставок"],
        ["Учебные запросы", "Выполнение запросов из варианта 11"],
        ["Добавление поставки", "INSERT в таблицу deliveries"],
        ["Удаление поставки", "DELETE выбранной записи из таблицы deliveries"],
        ["Изменение Rover1000", "UPDATE laptops SET ram_mb = 512"],
        ["Удаление SComp", "DELETE поставок SComp и записи поставщика"],
    ])

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Вывод", level=1)
    doc.add_paragraph(
        "В работе спроектирована база данных для учета поставок ноутбуков, построены ER-диаграмма и реляционная модель, "
        "подготовлены SQL-скрипты создания и заполнения таблиц, выполнены запросы варианта 11. "
        "Разработанное приложение подтверждает возможность чтения и изменения данных на сервере MySQL."
    )
    doc.save(OUT)


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    make_er_image()
    make_rel_image()
    schema = read(SQL / "01_schema.sql")
    seed = read(SQL / "02_seed.sql")
    queries = read(SQL / "03_queries.sql")
    build_markdown(schema, seed, queries)
    build_docx(schema, seed, queries)
    print(OUT)


if __name__ == "__main__":
    main()
